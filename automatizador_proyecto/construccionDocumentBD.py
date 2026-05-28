"""
generar_todo.py
===============
A partir de schema.json realiza tres tareas en secuencia:
  1. Crea la base de datos y sus tablas en PostgreSQL.
  2. Genera un diccionario de datos en formato Word (.docx).
  3. Genera un diccionario de datos en formato PDF.

Dependencias (100% Python):
    pip install psycopg[binary] python-docx fpdf2 --break-system-packages
"""

import json
import os
import sys
import datetime

# ──────────────────────────────────────────────
# CONFIGURACIÓN CENTRAL
# ──────────────────────────────────────────────
ARCHIVO_JSON = "schema.json"
ARCHIVO_DOCX = "Diccionario_Datos.docx"
ARCHIVO_PDF  = "Diccionario_Invernadero.pdf"

# Equivalencias entre tipos lógicos y tipos PostgreSQL
TIPOS_POSTGRES = {
    "identity": "BIGSERIAL",
    "string":   "VARCHAR",
    "decimal":  "DECIMAL",
    "integer":  "INTEGER",
    "boolean":  "BOOLEAN",
    "date":     "DATE",
    "datetime": "TIMESTAMP",
    "text":     "TEXT",
}

# Paleta de colores (RGB para python-docx)
COLOR_PRIMARIO = "1F5799"   # azul encabezado
COLOR_CABECERA = "D5E8F0"   # fondo fila de encabezado
COLOR_FILA_PAR = "EFF7FB"   # fondo fila par
COLOR_BLANCO   = "FFFFFF"


# ──────────────────────────────────────────────
# UTILIDADES COMUNES
# ──────────────────────────────────────────────

def cargar_modelo(ruta: str) -> dict:
    with open(ruta, "r", encoding="utf-8") as f:
        return json.load(f)


def tipo_postgres(attr: dict) -> str:
    tipo    = attr["type"]
    longitud = attr.get("length", 0)
    if tipo == "string":
        return f"VARCHAR({longitud})" if longitud else "VARCHAR"
    if tipo == "decimal":
        return f"DECIMAL({longitud},2)" if longitud else "DECIMAL"
    return TIPOS_POSTGRES.get(tipo, "TEXT")


# ──────────────────────────────────────────────
# MÓDULO 1 — CREACIÓN DE BASE DE DATOS
# ──────────────────────────────────────────────

def crear_base_de_datos(datos: dict) -> None:
    try:
        import psycopg
    except ImportError:
        print("⚠  psycopg no está instalado. Instala con: pip install psycopg[binary]")
        return

    host      = datos["host"]
    puerto    = int(datos["port"])
    usuario   = datos["username"]
    clave     = datos["password"]
    nombre_bd = datos["database"]
    entidades = datos["entities"]

    # Conectar al sistema para crear la BD si no existe
    with psycopg.connect(
        host=host, user=usuario, password=clave,
        dbname="postgres", port=puerto, autocommit=True
    ) as conn_sistema:
        with conn_sistema.cursor() as cur:
            cur.execute("SELECT 1 FROM pg_database WHERE datname = %s", (nombre_bd,))
            if not cur.fetchone():
                cur.execute(f"CREATE DATABASE {nombre_bd}")
                print(f"✓ Base de datos '{nombre_bd}' creada.")
            else:
                print(f"  La base de datos '{nombre_bd}' ya existe.")

    # Crear tablas
    with psycopg.connect(
        host=host, user=usuario, password=clave,
        dbname=nombre_bd, port=puerto
    ) as conn:
        with conn.cursor() as cur:
            for entidad in entidades:
                nombre_tabla = entidad["name"]
                columnas = []
                pks = []

                for attr in entidad["attributes"]:
                    col      = attr["name"]
                    tipo_sql = tipo_postgres(attr)
                    nulable  = attr.get("nullable", True)
                    es_pk    = attr.get("pk", False)

                    definicion = f"    {col} {tipo_sql}"
                    if not nulable:
                        definicion += " NOT NULL"
                    columnas.append(definicion)
                    if es_pk:
                        pks.append(col)

                if pks:
                    columnas.append(f"    PRIMARY KEY ({', '.join(pks)})")

                for restriccion in entidad.get("constrains", []):
                    if restriccion["type"] == "unique":
                        cols   = restriccion["columns"]
                        sufijo = "_".join(cols)
                        columnas.append(
                            f"    CONSTRAINT uq_{nombre_tabla}_{sufijo} UNIQUE ({', '.join(cols)})"
                        )

                sql_crear = (
                    f"CREATE TABLE IF NOT EXISTS {nombre_tabla} (\n"
                    + ",\n".join(columnas)
                    + "\n);"
                )
                print(f"\n{sql_crear}")
                cur.execute(sql_crear)

                for relacion in entidad.get("relations", []):
                    campo_local = relacion["local_field"]
                    tabla_ref   = relacion["entity"]
                    campo_ref   = relacion["related_field"]
                    nombre_fk   = f"fk_{nombre_tabla}_{campo_local}"
                    sql_fk = (
                        f"ALTER TABLE {nombre_tabla} "
                        f"ADD CONSTRAINT {nombre_fk} "
                        f"FOREIGN KEY ({campo_local}) "
                        f"REFERENCES {tabla_ref}({campo_ref});"
                    )
                    try:
                        cur.execute(sql_fk)
                        print(f"  → FK añadida: {nombre_fk}")
                    except Exception as err:
                        print(f"  ⚠  Error al añadir FK {nombre_fk}: {err}")

                print(f"✓ Tabla '{nombre_tabla}' lista.")

        conn.commit()

    print("\n✅ Base de datos y tablas generadas correctamente.\n")


# ──────────────────────────────────────────────
# MÓDULO 2 — GENERACIÓN DE DOCUMENTO WORD
# (100 % python-docx, sin Node.js)
# ──────────────────────────────────────────────

def generar_docx(datos: dict, ruta_salida: str) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("⚠  python-docx no está instalado. Instala con: pip install python-docx")
        return

    # ── helpers internos ──────────────────────────────────────────

    def hex_to_rgb(hex_str: str):
        h = hex_str.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def set_cell_bg(cell, hex_color: str):
        """Rellena el fondo de una celda con un color hex."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def set_cell_borders(cell):
        """Añade borde fino a todos los lados de la celda."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for lado in ("top", "left", "bottom", "right"):
            borde = OxmlElement(f"w:{lado}")
            borde.set(qn("w:val"),   "single")
            borde.set(qn("w:sz"),    "4")
            borde.set(qn("w:space"), "0")
            borde.set(qn("w:color"), "BFBFBF")
            tcBorders.append(borde)
        tcPr.append(tcBorders)

    def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        mar  = OxmlElement("w:tcMar")
        for lado, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
            m = OxmlElement(f"w:{lado}")
            m.set(qn("w:w"),    str(val))
            m.set(qn("w:type"), "dxa")
            mar.append(m)
        tcPr.append(mar)

    def celda_texto(cell, texto: str, negrita=False, fondo=COLOR_BLANCO,
                    alinear=WD_ALIGN_PARAGRAPH.LEFT, tamanio=9):
        """Escribe texto en una celda con estilo."""
        set_cell_bg(cell, fondo)
        set_cell_borders(cell)
        set_cell_margins(cell)
        para = cell.paragraphs[0]
        para.alignment = alinear
        run = para.add_run(str(texto or ""))
        run.bold      = negrita
        run.font.size = Pt(tamanio)
        run.font.name = "Arial"
        if negrita and fondo == COLOR_PRIMARIO:
            run.font.color.rgb = hex_to_rgb(COLOR_BLANCO)
        else:
            run.font.color.rgb = RGBColor(0, 0, 0)

    def parrafo_estilo(doc, texto, negrita=False, tamanio=10, color=None, espacio_antes=0, espacio_despues=6):
        p   = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(espacio_antes)
        p.paragraph_format.space_after  = Pt(espacio_despues)
        run = p.add_run(texto)
        run.bold      = negrita
        run.font.size = Pt(tamanio)
        run.font.name = "Arial"
        if color:
            run.font.color.rgb = hex_to_rgb(color)
        return p

    def tabla_atributos(doc, atributos):
        """Crea la tabla de atributos de una entidad."""
        cabeceras = ["Campo", "Tipo lógico", "Tipo PostgreSQL", "Long.", "Nulo", "PK", "Descripción"]
        anchos_cm = [3.5, 2.8, 3.5, 1.5, 1.5, 1.2, 5.5]  # suma ≈ 19.5 cm (A4 con márgenes 2.5cm)

        tabla = doc.add_table(rows=1, cols=len(cabeceras))
        tabla.style = "Table Grid"
        tabla.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Ajustar anchos de columna
        for i, ancho in enumerate(anchos_cm):
            tabla.columns[i].width = Cm(ancho)

        # Fila de encabezado
        fila_cab = tabla.rows[0]
        for i, cab in enumerate(cabeceras):
            celda_texto(fila_cab.cells[i], cab, negrita=True,
                        fondo=COLOR_PRIMARIO, tamanio=8)

        # Filas de datos
        for idx, attr in enumerate(atributos):
            fondo = COLOR_BLANCO if idx % 2 == 0 else COLOR_FILA_PAR
            fila  = tabla.add_row()
            vals  = [
                attr["name"],
                attr["type"],
                tipo_postgres(attr),
                str(attr.get("length") or "-"),
                "Sí" if attr.get("nullable", True) else "No",
                "Sí" if attr.get("pk", False) else "No",
                attr.get("description", ""),
            ]
            centrar = {3, 4, 5}
            for i, val in enumerate(vals):
                alinear = WD_ALIGN_PARAGRAPH.CENTER if i in centrar else WD_ALIGN_PARAGRAPH.LEFT
                celda_texto(fila.cells[i], val, fondo=fondo,
                            alinear=alinear, tamanio=8)
        return tabla

    def tabla_restricciones(doc, restricciones):
        cabeceras = ["Columna(s)", "Tipo", "Descripción"]
        anchos_cm = [4.5, 2.5, 12.5]

        tabla = doc.add_table(rows=1, cols=3)
        tabla.style = "Table Grid"
        fila_cab = tabla.rows[0]
        for i, cab in enumerate(cabeceras):
            celda_texto(fila_cab.cells[i], cab, negrita=True,
                        fondo=COLOR_CABECERA, tamanio=8)
        for i, ancho in enumerate(anchos_cm):
            tabla.columns[i].width = Cm(ancho)

        for idx, r in enumerate(restricciones):
            fondo = COLOR_BLANCO if idx % 2 == 0 else COLOR_FILA_PAR
            fila  = tabla.add_row()
            celda_texto(fila.cells[0], ", ".join(r["columns"]),   fondo=fondo, tamanio=8)
            celda_texto(fila.cells[1], r["type"].upper(),          fondo=fondo,
                        alinear=WD_ALIGN_PARAGRAPH.CENTER, tamanio=8)
            celda_texto(fila.cells[2], r.get("description", ""),  fondo=fondo, tamanio=8)
        return tabla

    def tabla_relaciones(doc, relaciones):
        cabeceras = ["Campo local", "Tabla ref.", "Campo ref.", "Cardinalidad", "Descripción"]
        anchos_cm = [3.5, 3.5, 3.5, 3.0, 6.0]

        tabla = doc.add_table(rows=1, cols=5)
        tabla.style = "Table Grid"
        fila_cab = tabla.rows[0]
        for i, cab in enumerate(cabeceras):
            celda_texto(fila_cab.cells[i], cab, negrita=True,
                        fondo=COLOR_CABECERA, tamanio=8)
        for i, ancho in enumerate(anchos_cm):
            tabla.columns[i].width = Cm(ancho)

        for idx, rel in enumerate(relaciones):
            fondo = COLOR_BLANCO if idx % 2 == 0 else COLOR_FILA_PAR
            fila  = tabla.add_row()
            celda_texto(fila.cells[0], rel["local_field"],          fondo=fondo, tamanio=8)
            celda_texto(fila.cells[1], rel["entity"],               fondo=fondo, tamanio=8)
            celda_texto(fila.cells[2], rel["related_field"],        fondo=fondo, tamanio=8)
            celda_texto(fila.cells[3], rel.get("type", ""),         fondo=fondo,
                        alinear=WD_ALIGN_PARAGRAPH.CENTER, tamanio=8)
            celda_texto(fila.cells[4], rel.get("description", ""),  fondo=fondo, tamanio=8)
        return tabla

    # ── Construcción del documento ────────────────────────────────

    doc = Document()

    # Márgenes de página (2.5 cm)
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # — Portada ───────────────────────────────────────────────────
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run(f"Diccionario de Datos – {datos['proyecto'].upper()}")
    run_titulo.bold           = True
    run_titulo.font.size      = Pt(20)
    run_titulo.font.name      = "Arial"
    run_titulo.font.color.rgb = hex_to_rgb(COLOR_PRIMARIO)
    titulo.paragraph_format.space_after = Pt(12)

    campos_portada = [
        ("Motor de base de datos", datos["motor"]),
        ("Servidor",               f"{datos['host']}:{datos['port']}"),
        ("Nombre de la base",      datos["database"]),
        ("Usuario de conexión",    datos["username"]),
        ("Fecha de generación",    datetime.datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]
    for etiq, val in campos_portada:
        p   = doc.add_paragraph()
        run = p.add_run(f"{etiq}: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run2 = p.add_run(val)
        run2.font.size = Pt(10)
        run2.font.name = "Arial"
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # — Índice ────────────────────────────────────────────────────
    parrafo_estilo(doc, "Índice de entidades", negrita=True, tamanio=14,
                   color=COLOR_PRIMARIO, espacio_antes=0, espacio_despues=8)

    for i, ent in enumerate(datos["entities"], 1):
        p   = doc.add_paragraph(style="List Number")
        run = p.add_run(f"{ent['name'].upper()} – {ent.get('description', '')}")
        run.font.size = Pt(10)
        run.font.name = "Arial"

    doc.add_page_break()

    # — Sección por entidad ───────────────────────────────────────
    for entidad in datos["entities"]:
        # Título de entidad
        parrafo_estilo(doc, f"Entidad: {entidad['name'].upper()}",
                       negrita=True, tamanio=14, color=COLOR_PRIMARIO,
                       espacio_antes=0, espacio_despues=4)

        if entidad.get("description"):
            parrafo_estilo(doc, entidad["description"], tamanio=10,
                           espacio_antes=0, espacio_despues=8)

        # Atributos
        parrafo_estilo(doc, "Atributos", negrita=True, tamanio=11,
                       color=COLOR_PRIMARIO, espacio_antes=4, espacio_despues=4)
        tabla_atributos(doc, entidad["attributes"])
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # Restricciones
        if entidad.get("constrains"):
            parrafo_estilo(doc, "Restricciones", negrita=True, tamanio=11,
                           color=COLOR_PRIMARIO, espacio_antes=4, espacio_despues=4)
            tabla_restricciones(doc, entidad["constrains"])
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # Relaciones
        if entidad.get("relations"):
            parrafo_estilo(doc, "Relaciones (claves foráneas)", negrita=True,
                           tamanio=11, color=COLOR_PRIMARIO,
                           espacio_antes=4, espacio_despues=4)
            tabla_relaciones(doc, entidad["relations"])
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

        doc.add_page_break()

    # — Leyenda ───────────────────────────────────────────────────
    parrafo_estilo(doc, "Leyenda", negrita=True, tamanio=14,
                   color=COLOR_PRIMARIO, espacio_antes=0, espacio_despues=8)

    entradas_leyenda = [
        ("PK",            "Clave primaria (Primary Key)"),
        ("FK",            "Clave foránea (Foreign Key)"),
        ("NULL / NO NULL","Indica si el campo admite valores nulos"),
        ("BIGSERIAL",     "Entero auto-incremental de 64 bits"),
        ("VARCHAR(n)",    "Cadena de caracteres de longitud variable hasta n"),
        ("DECIMAL(n,2)",  "Número decimal con n dígitos en total y 2 decimales"),
        ("TIMESTAMP",     "Fecha y hora"),
        ("BOOLEAN",       "Verdadero o falso"),
        ("one_to_one",    "Un registro se relaciona con exactamente uno"),
        ("many_to_one",   "Varios registros apuntan al mismo registro"),
    ]

    tabla_ley = doc.add_table(rows=1, cols=2)
    tabla_ley.style = "Table Grid"
    tabla_ley.columns[0].width = Cm(4)
    tabla_ley.columns[1].width = Cm(15.5)
    celda_texto(tabla_ley.rows[0].cells[0], "Término",     negrita=True, fondo=COLOR_CABECERA, tamanio=9)
    celda_texto(tabla_ley.rows[0].cells[1], "Significado", negrita=True, fondo=COLOR_CABECERA, tamanio=9)

    for idx, (term, sig) in enumerate(entradas_leyenda):
        fondo = COLOR_BLANCO if idx % 2 == 0 else COLOR_FILA_PAR
        fila  = tabla_ley.add_row()
        celda_texto(fila.cells[0], term, negrita=True, fondo=fondo, tamanio=9)
        celda_texto(fila.cells[1], sig,  fondo=fondo, tamanio=9)

    doc.save(ruta_salida)
    print(f"✅ Documento Word generado: {ruta_salida}")


# ──────────────────────────────────────────────
# MÓDULO 3 — GENERACIÓN DE PDF
# ──────────────────────────────────────────────

def generar_pdf(datos: dict, ruta_salida: str) -> None:
    try:
        from fpdf import FPDF
        from fpdf.enums import XPos, YPos
    except ImportError:
        print("⚠  fpdf2 no está instalado. Instala con: pip install fpdf2")
        return

    entidades = datos["entities"]

    AZUL      = (31,  87, 153)
    AZUL_CLAR = (213, 232, 240)
    GRIS_FILA = (239, 247, 251)
    BLANCO    = (255, 255, 255)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # — Portada ───────────────────────────────────────────────────
    pdf.add_page()
    pdf.set_font("helvetica", "B", 22)
    pdf.set_text_color(*AZUL)
    pdf.cell(0, 14, f"Diccionario de Datos: {datos['proyecto'].upper()}",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.ln(6)

    pdf.set_font("helvetica", size=10)
    pdf.set_text_color(60, 60, 60)
    campos_portada = [
        ("Motor de base de datos", datos["motor"]),
        ("Servidor",               f"{datos['host']}:{datos['port']}"),
        ("Nombre de la base",      datos["database"]),
        ("Usuario de conexion",    datos["username"]),
        ("Fecha de generacion",    datetime.datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]
    for etiqueta, valor in campos_portada:
        pdf.set_font("helvetica", "B", 10)
        pdf.cell(60, 7, f"{etiqueta}:", new_x=XPos.RIGHT, new_y=YPos.TOP)
        pdf.set_font("helvetica", size=10)
        pdf.cell(0, 7, valor, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(8)

    # — Índice ────────────────────────────────────────────────────
    pdf.set_font("helvetica", "B", 14)
    pdf.set_text_color(*AZUL)
    pdf.cell(0, 10, "Indice de entidades", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_text_color(0, 0, 0)
    for i, entidad in enumerate(entidades, 1):
        pdf.set_font("helvetica", "B", 10)
        pdf.cell(10, 7, f"{i}.", new_x=XPos.RIGHT, new_y=YPos.TOP)
        pdf.set_font("helvetica", size=10)
        pdf.cell(0, 7, f"{entidad['name'].upper()} - {entidad.get('description', '')}",
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # — Entidades ─────────────────────────────────────────────────
    for entidad in entidades:
        pdf.add_page()

        pdf.set_font("helvetica", "B", 16)
        pdf.set_text_color(*AZUL)
        pdf.cell(0, 11, f"Entidad: {entidad['name'].upper()}",
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        pdf.set_font("helvetica", "I", 10)
        pdf.set_text_color(60, 60, 60)
        pdf.multi_cell(0, 5, entidad.get("description", ""),
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(3)

        # Tabla de atributos
        pdf.set_font("helvetica", "B", 10)
        pdf.set_text_color(*AZUL)
        pdf.cell(0, 8, "Atributos", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)

        anchos_attr = [28, 20, 28, 12, 12, 10, 80]
        cabeceras   = ["Campo", "Tipo logico", "Tipo PgSQL", "Long", "Nulo", "PK", "Descripcion"]
        pdf.set_font("helvetica", "B", 8)
        pdf.set_fill_color(*AZUL_CLAR)
        for cab, ancho in zip(cabeceras, anchos_attr):
            pdf.cell(ancho, 7, cab, border=1, fill=True)
        pdf.ln()

        pdf.set_font("helvetica", size=7)
        for idx, attr in enumerate(entidad["attributes"]):
            fondo = GRIS_FILA if idx % 2 else BLANCO
            pdf.set_fill_color(*fondo)
            valores = [
                attr["name"],
                attr["type"],
                tipo_postgres(attr),
                str(attr.get("length") or "-"),
                "Si" if attr.get("nullable", True) else "No",
                "Si" if attr.get("pk", False) else "No",
                attr.get("description", ""),
            ]
            for val, ancho in zip(valores, anchos_attr):
                pdf.cell(ancho, 6, str(val)[:40], border=1, fill=True)
            pdf.ln()
        pdf.ln(4)

        # Restricciones
        if entidad.get("constrains"):
            pdf.set_font("helvetica", "B", 10)
            pdf.set_text_color(*AZUL)
            pdf.cell(0, 8, "Restricciones", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("helvetica", size=8)
            for restriccion in entidad["constrains"]:
                cols = ", ".join(restriccion["columns"])
                pdf.multi_cell(
                    0, 5,
                    f"- [{restriccion['type'].upper()}] ({cols}): {restriccion.get('description', '')}",
                    new_x=XPos.LMARGIN, new_y=YPos.NEXT,
                )
            pdf.ln(3)

        # Relaciones
        if entidad.get("relations"):
            pdf.set_font("helvetica", "B", 10)
            pdf.set_text_color(*AZUL)
            pdf.cell(0, 8, "Relaciones (claves foraneas)", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("helvetica", size=8)
            for rel in entidad["relations"]:
                pdf.multi_cell(
                    0, 5,
                    f"- {rel['local_field']} -> {rel['entity']}({rel['related_field']}) "
                    f"[{rel.get('type', '')}]: {rel.get('description', '')}",
                    new_x=XPos.LMARGIN, new_y=YPos.NEXT,
                )
            pdf.ln(3)

    # — Leyenda ───────────────────────────────────────────────────
    pdf.add_page()
    pdf.set_font("helvetica", "B", 14)
    pdf.set_text_color(*AZUL)
    pdf.cell(0, 10, "Leyenda", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    entradas_leyenda = [
        ("PK",            "Clave primaria (Primary Key)"),
        ("FK",            "Clave foranea (Foreign Key)"),
        ("NULL / NO NULL","Indica si el campo admite valores nulos"),
        ("BIGSERIAL",     "Entero auto-incremental de 64 bits"),
        ("VARCHAR(n)",    "Cadena de caracteres variable hasta n posiciones"),
        ("DECIMAL(n,2)",  "Numero decimal con n digitos totales y 2 decimales"),
        ("TIMESTAMP",     "Fecha y hora (fecha + tiempo)"),
        ("BOOLEAN",       "Valor logico: verdadero o falso"),
        ("one_to_one",    "Un registro se relaciona con exactamente uno"),
        ("many_to_one",   "Varios registros apuntan al mismo registro"),
    ]
    for termino, significado in entradas_leyenda:
        pdf.set_font("helvetica", "B", 9)
        pdf.cell(38, 7, termino, new_x=XPos.RIGHT, new_y=YPos.TOP)
        pdf.set_font("helvetica", size=9)
        pdf.cell(0, 7, significado, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.ln(8)
    pdf.set_font("helvetica", "I", 8)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 5, f"Generado automaticamente desde {ARCHIVO_JSON}",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")

    pdf.output(ruta_salida)
    print(f"✅ Documento PDF generado: {ruta_salida}")


# ──────────────────────────────────────────────
# PUNTO DE ENTRADA
# ──────────────────────────────────────────────

def main():
    print("=" * 50)
    print("  Generador de diccionario de datos")
    print("=" * 50)

    if not os.path.exists(ARCHIVO_JSON):
        print(f"❌ No se encontró el archivo de modelo: {ARCHIVO_JSON}")
        sys.exit(1)

    modelo = cargar_modelo(ARCHIVO_JSON)

    print("\n[1/3] Creando base de datos y tablas…")
    crear_base_de_datos(modelo)

    print("\n[2/3] Generando documento Word…")
    generar_docx(modelo, ARCHIVO_DOCX)

    print("\n[3/3] Generando documento PDF…")
    generar_pdf(modelo, ARCHIVO_PDF)

    print("\n" + "=" * 50)
    print("  Proceso finalizado.")
    print(f"  Word : {ARCHIVO_DOCX}")
    print(f"  PDF  : {ARCHIVO_PDF}")
    print("=" * 50)


if __name__ == "__main__":
    main()